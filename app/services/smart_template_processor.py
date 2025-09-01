"""
Intelligent Template Analysis and Upload System
Implements OCR, coordinate mapping, visual placeholder selection, and 
smart suggestion system for advanced template creation capabilities.
"""

import os
import json
import uuid
import asyncio
from typing import Dict, List, Optional, Any, Tuple, Set
from dataclasses import dataclass
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import cv2
import numpy as np
from io import BytesIO
import logging

from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import PyPDF2
from pdf2image import convert_from_path

from app.models.template import Template, Placeholder
from config import settings

# Configure logging
template_logger = logging.getLogger('smart_template')

@dataclass
class TextInstance:
    """Represents a text instance found in document"""
    text: str
    position: Dict[str, Any]
    formatting: Dict[str, Any]
    page_number: int
    confidence: float
    context: str

@dataclass
class PlaceholderSuggestion:
    """Suggested placeholder based on text analysis"""
    name: str
    display_name: str
    instances: List[TextInstance]
    similarity_score: float
    semantic_type: str
    consolidation_opportunities: List[str]

class SmartTemplateProcessor:
    """
    Intelligent template processing with OCR, coordinate mapping,
    and visual placeholder selection capabilities
    """
    
    def __init__(self):
        self.ocr_config = '--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?@#$%^&*()_+-=[]{}|;:,.<>?/~`\'" '
        
    async def process_template_upload(
        self,
        file_content: bytes,
        filename: str,
        user_id: int,
        db: Session
    ) -> Dict[str, Any]:
        """
        Complete template processing workflow with intelligent analysis
        """
        processing_start = asyncio.get_event_loop().time()
        
        try:
            # Determine file type and extract text
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension == '.pdf':
                text_instances = await self._process_pdf_document(file_content)
            elif file_extension in ['.docx', '.doc']:
                text_instances = await self._process_word_document(file_content)
            elif file_extension in ['.png', '.jpg', '.jpeg']:
                text_instances = await self._process_image_document(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Create searchable text index
            text_index = self._create_searchable_index(text_instances)
            
            # Generate placeholder suggestions
            placeholder_suggestions = await self._generate_placeholder_suggestions(text_instances)
            
            # Calculate processing stats
            processing_time = (asyncio.get_event_loop().time() - processing_start) * 1000
            
            result = {
                'status': 'success',
                'text_instances_found': len(text_instances),
                'searchable_index': text_index,
                'placeholder_suggestions': [
                    {
                        'name': s.name,
                        'display_name': s.display_name,
                        'instance_count': len(s.instances),
                        'similarity_score': s.similarity_score,
                        'semantic_type': s.semantic_type,
                        'consolidation_opportunities': s.consolidation_opportunities
                    }
                    for s in placeholder_suggestions
                ],
                'processing_time_ms': processing_time,
                'template_ready_for_creation': len(placeholder_suggestions) > 0
            }
            
            template_logger.info(
                f"Template processed: {len(text_instances)} text instances, "
                f"{len(placeholder_suggestions)} placeholder suggestions in {processing_time:.2f}ms"
            )
            
            return result
            
        except Exception as e:
            template_logger.error(f"Template processing failed: {e}")
            return {
                'status': 'error',
                'error': str(e),
                'processing_time_ms': (asyncio.get_event_loop().time() - processing_start) * 1000
            }
    
    async def _process_pdf_document(self, file_content: bytes) -> List[TextInstance]:
        """
        Process PDF document with OCR and coordinate extraction
        """
        text_instances = []
        
        try:
            # Convert PDF to images
            pdf_file = BytesIO(file_content)
            images = convert_from_path(pdf_file, dpi=300)
            
            for page_num, image in enumerate(images):
                # Convert PIL Image to OpenCV format
                opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                
                # Extract text with coordinates
                ocr_data = pytesseract.image_to_data(
                    opencv_image, 
                    output_type=pytesseract.Output.DICT,
                    config=self.ocr_config
                )
                
                # Process OCR results
                for i, text in enumerate(ocr_data['text']):
                    confidence = ocr_data['conf'][i]
                    
                    if confidence > 50 and text.strip():  # Filter low confidence
                        text_instance = TextInstance(
                            text=text.strip(),
                            position={
                                'page': page_num + 1,
                                'x': ocr_data['left'][i],
                                'y': ocr_data['top'][i],
                                'width': ocr_data['width'][i],
                                'height': ocr_data['height'][i]
                            },
                            formatting={
                                'font_size': self._estimate_font_size(ocr_data['height'][i]),
                                'bold': self._detect_bold_text(opencv_image, ocr_data, i),
                                'italic': False  # Basic implementation
                            },
                            page_number=page_num + 1,
                            confidence=confidence,
                            context=self._determine_text_context(ocr_data, i)
                        )
                        text_instances.append(text_instance)
            
        except Exception as e:
            template_logger.error(f"PDF processing failed: {e}")
        
        return text_instances
    
    async def _process_word_document(self, file_content: bytes) -> List[TextInstance]:
        """
        Process Word document with precise coordinate extraction
        """
        text_instances = []
        
        try:
            doc = DocxDocument(BytesIO(file_content))
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Extract words and their formatting
                    for run_idx, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            text_instance = TextInstance(
                                text=run.text.strip(),
                                position={
                                    'paragraph_index': para_idx,
                                    'run_index': run_idx,
                                    'start_index': 0,
                                    'end_index': len(run.text)
                                },
                                formatting={
                                    'font_size': run.font.size.pt if run.font.size else 12,
                                    'font_family': run.font.name or 'Times New Roman',
                                    'bold': run.bold or False,
                                    'italic': run.italic or False,
                                    'underline': run.underline or False
                                },
                                page_number=1,  # Simplified for Word docs
                                confidence=100.0,  # High confidence for Word
                                context=self._determine_word_context(para_idx, len(doc.paragraphs))
                            )
                            text_instances.append(text_instance)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if paragraph.text.strip():
                                text_instance = TextInstance(
                                    text=paragraph.text.strip(),
                                    position={
                                        'table_index': table_idx,
                                        'row_index': row_idx,
                                        'cell_index': cell_idx,
                                        'paragraph_index': para_idx
                                    },
                                    formatting={
                                        'font_size': 12,  # Default for tables
                                        'bold': False,
                                        'italic': False
                                    },
                                    page_number=1,
                                    confidence=100.0,
                                    context='table'
                                )
                                text_instances.append(text_instance)
                                
        except Exception as e:
            template_logger.error(f"Word document processing failed: {e}")
        
        return text_instances
    
    async def _process_image_document(self, file_content: bytes) -> List[TextInstance]:
        """
        Process image document with advanced OCR
        """
        text_instances = []
        
        try:
            # Load image
            image = Image.open(BytesIO(file_content))
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Preprocess image for better OCR
            processed_image = self._preprocess_image_for_ocr(opencv_image)
            
            # Extract text with coordinates
            ocr_data = pytesseract.image_to_data(
                processed_image,
                output_type=pytesseract.Output.DICT,
                config=self.ocr_config
            )
            
            # Process OCR results
            for i, text in enumerate(ocr_data['text']):
                confidence = ocr_data['conf'][i]
                
                if confidence > 40 and text.strip():  # Lower threshold for images
                    text_instance = TextInstance(
                        text=text.strip(),
                        position={
                            'x': ocr_data['left'][i],
                            'y': ocr_data['top'][i],
                            'width': ocr_data['width'][i],
                            'height': ocr_data['height'][i]
                        },
                        formatting={
                            'font_size': self._estimate_font_size(ocr_data['height'][i]),
                            'bold': False,  # Difficult to detect in images
                            'italic': False
                        },
                        page_number=1,
                        confidence=confidence,
                        context=self._determine_image_text_context(ocr_data, i)
                    )
                    text_instances.append(text_instance)
                    
        except Exception as e:
            template_logger.error(f"Image processing failed: {e}")
        
        return text_instances
    
    def _preprocess_image_for_ocr(self, image: np.ndarray) -> np.ndarray:
        """
        Preprocess image to improve OCR accuracy
        """
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        
        # Apply threshold to get binary image
        _, binary = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Morphological operations to clean up
        kernel = np.ones((1, 1), np.uint8)
        cleaned = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
        
        return cleaned
    
    def _create_searchable_index(self, text_instances: List[TextInstance]) -> Dict[str, Any]:
        """
        Create searchable text index for intelligent suggestions
        """
        index = {}
        
        for instance in text_instances:
            words = instance.text.split()
            
            for word in words:
                word_clean = word.lower().strip('.,!?;:"()[]{}')
                
                if len(word_clean) > 2:  # Skip very short words
                    if word_clean not in index:
                        index[word_clean] = []
                    
                    index[word_clean].append({
                        'original_text': instance.text,
                        'position': instance.position,
                        'formatting': instance.formatting,
                        'page_number': instance.page_number,
                        'confidence': instance.confidence,
                        'context': instance.context
                    })
        
        return index
    
    async def _generate_placeholder_suggestions(
        self,
        text_instances: List[TextInstance]
    ) -> List[PlaceholderSuggestion]:
        """
        Generate intelligent placeholder suggestions based on text analysis
        """
        suggestions = []
        
        # Group similar text instances
        text_groups = self._group_similar_texts(text_instances)
        
        # Analyze each group for placeholder potential
        for group_text, instances in text_groups.items():
            if len(instances) > 1 or self._is_placeholder_candidate(group_text):
                suggestion = PlaceholderSuggestion(
                    name=self._generate_placeholder_name(group_text),
                    display_name=self._generate_display_name(group_text),
                    instances=instances,
                    similarity_score=self._calculate_similarity_score(instances),
                    semantic_type=self._classify_semantic_type(group_text),
                    consolidation_opportunities=self._find_consolidation_opportunities(
                        group_text, text_groups
                    )
                )
                suggestions.append(suggestion)
        
        # Sort by relevance
        suggestions.sort(key=lambda x: x.similarity_score, reverse=True)
        
        return suggestions[:20]  # Return top 20 suggestions
    
    def _group_similar_texts(self, text_instances: List[TextInstance]) -> Dict[str, List[TextInstance]]:
        """
        Group text instances that are likely the same placeholder
        """
        groups = {}
        
        for instance in text_instances:
            text = instance.text.lower().strip()
            
            # Find existing similar group
            matched_group = None
            for group_key in groups.keys():
                if self._texts_are_similar(text, group_key, 0.8):
                    matched_group = group_key
                    break
            
            if matched_group:
                groups[matched_group].append(instance)
            else:
                groups[text] = [instance]
        
        return groups
    
    def _texts_are_similar(self, text1: str, text2: str, threshold: float) -> bool:
        """
        Determine if two texts are similar enough to be the same placeholder
        """
        # Exact match
        if text1 == text2:
            return True
        
        # Length similarity
        len_diff = abs(len(text1) - len(text2)) / max(len(text1), len(text2))
        if len_diff > 0.5:
            return False
        
        # Character similarity (simple implementation)
        common_chars = sum(1 for a, b in zip(text1, text2) if a == b)
        similarity = common_chars / max(len(text1), len(text2))
        
        return similarity >= threshold
    
    def _is_placeholder_candidate(self, text: str) -> bool:
        """
        Determine if text is a good placeholder candidate
        """
        # Common placeholder patterns
        placeholder_indicators = [
            'name', 'date', 'address', 'email', 'phone',
            'signature', 'title', 'company', 'position',
            'amount', 'number', 'reference', 'id'
        ]
        
        text_lower = text.lower()
        
        # Check if it contains placeholder indicators
        for indicator in placeholder_indicators:
            if indicator in text_lower:
                return True
        
        # Check if it looks like sample data
        if any(char.isdigit() for char in text) and len(text) > 3:
            return True
        
        # Check for email patterns
        if '@' in text and '.' in text:
            return True
        
        # Check for phone patterns
        if any(char.isdigit() for char in text) and len([c for c in text if c.isdigit()]) >= 7:
            return True
        
        return False
    
    def _generate_placeholder_name(self, text: str) -> str:
        """
        Generate appropriate placeholder name from text
        """
        text_lower = text.lower().strip()
        
        # Predefined mappings
        name_mappings = {
            'name': 'full_name',
            'email': 'email_address',
            'phone': 'phone_number',
            'address': 'address',
            'date': 'date',
            'signature': 'signature',
            'company': 'company_name',
            'title': 'title'
        }
        
        # Check for exact matches
        for key, value in name_mappings.items():
            if key in text_lower:
                return value
        
        # Generate from text
        words = text_lower.split()
        if len(words) <= 3:
            return '_'.join(words)
        else:
            return f"{words[0]}_{words[-1]}"
    
    def _generate_display_name(self, text: str) -> str:
        """
        Generate user-friendly display name
        """
        display_mappings = {
            'full_name': 'Full Name',
            'email_address': 'Email Address',
            'phone_number': 'Phone Number',
            'address': 'Address',
            'date': 'Date',
            'signature': 'Signature',
            'company_name': 'Company Name'
        }
        
        placeholder_name = self._generate_placeholder_name(text)
        
        if placeholder_name in display_mappings:
            return display_mappings[placeholder_name]
        
        return placeholder_name.replace('_', ' ').title()
    
    def _calculate_similarity_score(self, instances: List[TextInstance]) -> float:
        """
        Calculate similarity score for placeholder suggestion
        """
        base_score = min(len(instances) * 0.3, 1.0)  # More instances = higher score
        
        # Bonus for high confidence OCR results
        avg_confidence = sum(i.confidence for i in instances) / len(instances)
        confidence_bonus = (avg_confidence / 100) * 0.3
        
        # Bonus for consistent formatting
        formatting_consistency = self._calculate_formatting_consistency(instances)
        formatting_bonus = formatting_consistency * 0.2
        
        # Bonus for semantic relevance
        semantic_bonus = 0.2 if self._is_placeholder_candidate(instances[0].text) else 0
        
        total_score = base_score + confidence_bonus + formatting_bonus + semantic_bonus
        return min(total_score, 1.0)
    
    def _calculate_formatting_consistency(self, instances: List[TextInstance]) -> float:
        """
        Calculate how consistent formatting is across instances
        """
        if len(instances) < 2:
            return 1.0
        
        first_format = instances[0].formatting
        consistent_properties = 0
        total_properties = 0
        
        for prop in ['font_size', 'bold', 'italic']:
            total_properties += 1
            if all(i.formatting.get(prop) == first_format.get(prop) for i in instances):
                consistent_properties += 1
        
        return consistent_properties / total_properties if total_properties > 0 else 0.0
    
    def _classify_semantic_type(self, text: str) -> str:
        """
        Classify text into semantic type
        """
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['name', 'applicant', 'student']):
            return 'name'
        elif any(word in text_lower for word in ['address', 'location', 'residence']):
            return 'address'
        elif any(word in text_lower for word in ['date', 'birth', 'issued']):
            return 'date'
        elif any(word in text_lower for word in ['signature', 'sign']):
            return 'signature'
        elif '@' in text and '.' in text:
            return 'email'
        elif any(char.isdigit() for char in text) and len([c for c in text if c.isdigit()]) >= 7:
            return 'phone'
        else:
            return 'text'
    
    def _find_consolidation_opportunities(
        self,
        current_text: str,
        all_groups: Dict[str, List[TextInstance]]
    ) -> List[str]:
        """
        Find other text groups that could be consolidated with current group
        """
        opportunities = []
        current_semantic = self._classify_semantic_type(current_text)
        
        for other_text in all_groups.keys():
            if other_text != current_text:
                other_semantic = self._classify_semantic_type(other_text)
                
                if current_semantic == other_semantic and current_semantic != 'text':
                    opportunities.append(other_text)
        
        return opportunities[:5]  # Limit to top 5 opportunities
    
    def _estimate_font_size(self, height: int) -> int:
        """
        Estimate font size from text height
        """
        # Rough estimation based on typical font metrics
        return max(8, min(72, int(height * 0.75)))
    
    def _detect_bold_text(self, image: np.ndarray, ocr_data: Dict, index: int) -> bool:
        """
        Attempt to detect if text is bold (basic implementation)
        """
        try:
            x, y, w, h = ocr_data['left'][index], ocr_data['top'][index], ocr_data['width'][index], ocr_data['height'][index]
            text_region = image[y:y+h, x:x+w]
            
            # Calculate average line thickness (simplified approach)
            gray = cv2.cvtColor(text_region, cv2.COLOR_BGR2GRAY)
            _, binary = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY_INV)
            
            # Count black pixels (text)
            text_pixels = np.sum(binary == 255)
            total_pixels = binary.shape[0] * binary.shape[1]
            
            # If more than 25% is text, consider it bold
            return (text_pixels / total_pixels) > 0.25
        except:
            return False
    
    def _determine_text_context(self, ocr_data: Dict, index: int) -> str:
        """
        Determine context of text (header, body, footer)
        """
        y_position = ocr_data['top'][index]
        image_height = max(ocr_data['top']) + max(ocr_data['height'])
        
        relative_position = y_position / image_height
        
        if relative_position < 0.2:
            return 'header'
        elif relative_position > 0.8:
            return 'footer'
        else:
            return 'body'
    
    def _determine_word_context(self, paragraph_index: int, total_paragraphs: int) -> str:
        """
        Determine context for Word document text
        """
        relative_position = paragraph_index / max(total_paragraphs, 1)
        
        if relative_position < 0.15:
            return 'header'
        elif relative_position > 0.85:
            return 'footer'
        else:
            return 'body'
    
    def _determine_image_text_context(self, ocr_data: Dict, index: int) -> str:
        """
        Determine context for image text
        """
        return self._determine_text_context(ocr_data, index)

# Global processor instance
smart_template_processor = SmartTemplateProcessor()