"""
Advanced Signature Canvas Processing Service
Handles signature capture, background removal, auto-sizing, and seamless 
document integration with intelligent placement and quality enhancement.
"""

import os
import io
import base64
import uuid
from typing import Optional, Dict, Any, Tuple, List
from PIL import Image, ImageOps, ImageFilter, ImageEnhance
import numpy as np
from io import BytesIO
import logging

from sqlalchemy.orm import Session
from app.models.signature import Signature, SignatureType, SignatureStatus
from app.models.document import Document
from config import settings

# Configure logging
signature_logger = logging.getLogger('signature_processing')

class SignatureProcessingOptions:
    """Configuration options for signature processing"""
    
    def __init__(self):
        # Background removal settings
        self.remove_background = True
        self.background_threshold = 240  # RGB values above this are considered background
        self.edge_smoothing = True
        
        # Enhancement settings
        self.enhance_contrast = True
        self.enhance_sharpness = True
        self.line_thickness_enhancement = True
        
        # Size optimization
        self.auto_resize = True
        self.target_width = 300
        self.target_height = 100
        self.maintain_aspect_ratio = True
        
        # Quality settings
        self.output_format = 'PNG'
        self.compression_quality = 95
        self.dpi = 300

class AdvancedSignatureProcessor:
    """
    Advanced signature processing with AI-powered background removal,
    quality enhancement, and intelligent document placement
    """
    
    def __init__(self):
        self.processing_options = SignatureProcessingOptions()
        
    async def process_signature_canvas(
        self,
        signature_data: str,
        processing_options: Optional[SignatureProcessingOptions] = None
    ) -> Tuple[bytes, Dict[str, Any]]:
        """
        Process signature from canvas data with advanced enhancement
        
        Args:
            signature_data: Base64 encoded canvas image data
            processing_options: Custom processing options
            
        Returns:
            Tuple of (processed_image_bytes, processing_stats)
        """
        if not signature_data:
            raise ValueError("Signature data is required")
        
        options = processing_options or self.processing_options
        processing_stats = {
            'original_size': 0,
            'processed_size': 0,
            'background_removed': False,
            'enhanced': False,
            'resized': False,
            'processing_time_ms': 0
        }
        
        import time
        start_time = time.time()
        
        try:
            # Decode canvas data
            image = self._decode_canvas_data(signature_data)
            processing_stats['original_size'] = len(signature_data)
            
            # Remove background
            if options.remove_background:
                image = self._remove_background_intelligent(image, options)
                processing_stats['background_removed'] = True
            
            # Enhance signature quality
            if options.enhance_contrast or options.enhance_sharpness:
                image = self._enhance_signature_quality(image, options)
                processing_stats['enhanced'] = True
            
            # Resize and optimize
            if options.auto_resize:
                image = self._resize_signature_optimal(image, options)
                processing_stats['resized'] = True
            
            # Convert to output format
            processed_bytes = self._convert_to_output_format(image, options)
            processing_stats['processed_size'] = len(processed_bytes)
            
            processing_time = (time.time() - start_time) * 1000
            processing_stats['processing_time_ms'] = processing_time
            
            signature_logger.info(
                f"Signature processed: {processing_stats['original_size']} -> "
                f"{processing_stats['processed_size']} bytes in {processing_time:.2f}ms"
            )
            
            return processed_bytes, processing_stats
            
        except Exception as e:
            signature_logger.error(f"Signature processing failed: {e}")
            raise
    
    def _decode_canvas_data(self, signature_data: str) -> Image.Image:
        """
        Decode base64 canvas data to PIL Image
        """
        try:
            # Remove data URL prefix if present
            if signature_data.startswith('data:image'):
                signature_data = signature_data.split(',')[1]
            
            # Decode base64
            image_bytes = base64.b64decode(signature_data)
            
            # Open as PIL Image
            image = Image.open(BytesIO(image_bytes))
            
            # Convert to RGBA for processing
            if image.mode != 'RGBA':
                image = image.convert('RGBA')
            
            return image
            
        except Exception as e:
            raise ValueError(f"Failed to decode signature data: {e}")
    
    def _remove_background_intelligent(
        self,
        image: Image.Image,
        options: SignatureProcessingOptions
    ) -> Image.Image:
        """
        Intelligent background removal using multiple techniques
        """
        # Convert to numpy array for processing
        img_array = np.array(image)
        
        # Method 1: Threshold-based removal (for white/light backgrounds)
        img_processed = self._remove_background_threshold(img_array, options.background_threshold)
        
        # Method 2: Edge-based enhancement (preserve signature strokes)
        if options.edge_smoothing:
            img_processed = self._enhance_edges(img_processed)
        
        # Method 3: Remove isolated noise pixels
        img_processed = self._remove_noise_pixels(img_processed)
        
        return Image.fromarray(img_processed, 'RGBA')
    
    def _remove_background_threshold(self, img_array: np.ndarray, threshold: int) -> np.ndarray:
        """
        Remove background using color threshold
        """
        # Create alpha channel based on brightness
        brightness = np.mean(img_array[:, :, :3], axis=2)
        
        # Set transparent pixels where brightness is above threshold
        alpha_channel = np.where(brightness >= threshold, 0, 255).astype(np.uint8)
        
        # Preserve original RGB, update alpha
        result = img_array.copy()
        result[:, :, 3] = alpha_channel
        
        return result
    
    def _enhance_edges(self, img_array: np.ndarray) -> np.ndarray:
        """
        Enhance signature edges while maintaining smoothness
        """
        # Convert to PIL for filtering
        img = Image.fromarray(img_array, 'RGBA')
        
        # Apply edge enhancement
        img = img.filter(ImageFilter.EDGE_ENHANCE_MORE)
        
        # Smooth sharp edges
        img = img.filter(ImageFilter.SMOOTH_MORE)
        
        return np.array(img)
    
    def _remove_noise_pixels(self, img_array: np.ndarray) -> np.ndarray:
        """
        Remove isolated noise pixels from signature
        """
        # Simple morphological operations to remove small artifacts
        from scipy import ndimage
        
        # Create binary mask from alpha channel
        alpha = img_array[:, :, 3]
        binary_mask = alpha > 0
        
        # Remove small objects (noise)
        cleaned_mask = ndimage.binary_opening(binary_mask, structure=np.ones((3, 3)))
        
        # Apply cleaned mask
        result = img_array.copy()
        result[:, :, 3] = np.where(cleaned_mask, alpha, 0)
        
        return result
    
    def _enhance_signature_quality(
        self,
        image: Image.Image,
        options: SignatureProcessingOptions
    ) -> Image.Image:
        """
        Enhance signature quality with contrast and sharpness adjustments
        """
        enhanced_image = image
        
        # Enhance contrast
        if options.enhance_contrast:
            enhancer = ImageEnhance.Contrast(enhanced_image)
            enhanced_image = enhancer.enhance(1.3)  # Increase contrast by 30%
        
        # Enhance sharpness
        if options.enhance_sharpness:
            enhancer = ImageEnhance.Sharpness(enhanced_image)
            enhanced_image = enhancer.enhance(1.2)  # Increase sharpness by 20%
        
        # Enhance line thickness if needed
        if options.line_thickness_enhancement:
            enhanced_image = self._enhance_line_thickness(enhanced_image)
        
        return enhanced_image
    
    def _enhance_line_thickness(self, image: Image.Image) -> Image.Image:
        """
        Enhance thin signature lines for better visibility
        """
        # Apply dilation to thicken lines slightly
        from PIL import ImageFilter
        
        # Custom kernel for line thickening
        kernel = ImageFilter.Kernel(
            (3, 3),
            [1, 1, 1,
             1, 2, 1,
             1, 1, 1],
            scale=10
        )
        
        return image.filter(kernel)
    
    def _resize_signature_optimal(
        self,
        image: Image.Image,
        options: SignatureProcessingOptions
    ) -> Image.Image:
        """
        Resize signature with optimal quality preservation
        """
        original_width, original_height = image.size
        target_width = options.target_width
        target_height = options.target_height
        
        if options.maintain_aspect_ratio:
            # Calculate aspect ratio
            aspect_ratio = original_width / original_height
            
            # Determine optimal dimensions
            if aspect_ratio > (target_width / target_height):
                # Width is the limiting factor
                new_width = target_width
                new_height = int(target_width / aspect_ratio)
            else:
                # Height is the limiting factor
                new_height = target_height
                new_width = int(target_height * aspect_ratio)
        else:
            new_width = target_width
            new_height = target_height
        
        # Use high-quality resampling
        resized_image = image.resize(
            (new_width, new_height),
            Image.Resampling.LANCZOS
        )
        
        return resized_image
    
    def _convert_to_output_format(
        self,
        image: Image.Image,
        options: SignatureProcessingOptions
    ) -> bytes:
        """
        Convert processed image to desired output format
        """
        output_buffer = BytesIO()
        
        # Set DPI for high-quality output
        dpi = (options.dpi, options.dpi)
        
        if options.output_format.upper() == 'PNG':
            image.save(
                output_buffer,
                format='PNG',
                dpi=dpi,
                optimize=True
            )
        elif options.output_format.upper() == 'JPEG':
            # Convert RGBA to RGB for JPEG
            if image.mode == 'RGBA':
                # Create white background
                background = Image.new('RGB', image.size, (255, 255, 255))
                background.paste(image, mask=image.split()[3])  # Use alpha as mask
                image = background
            
            image.save(
                output_buffer,
                format='JPEG',
                quality=options.compression_quality,
                dpi=dpi,
                optimize=True
            )
        
        return output_buffer.getvalue()
    
    async def save_signature_to_database(
        self,
        db: Session,
        signature_data: bytes,
        user_id: int,
        document_id: Optional[int] = None,
        signature_type: SignatureType = SignatureType.CANVAS,
        processing_stats: Optional[Dict[str, Any]] = None
    ) -> Signature:
        """
        Save processed signature to database
        """
        # Generate unique filename
        signature_filename = f"signature_{uuid.uuid4().hex}.png"
        signature_path = os.path.join(settings.SIGNATURES_PATH, signature_filename)
        
        # Save signature file
        os.makedirs(settings.SIGNATURES_PATH, exist_ok=True)
        with open(signature_path, 'wb') as f:
            f.write(signature_data)
        
        # Create database record
        signature = Signature(
            user_id=user_id,
            document_id=document_id,
            type=signature_type,
            file_path=signature_filename,
            file_size=len(signature_data),
            processing_metadata=processing_stats or {},
            status=SignatureStatus.PROCESSED
        )
        
        db.add(signature)
        db.commit()
        db.refresh(signature)
        
        return signature
    
    async def integrate_signature_with_document(
        self,
        document_path: str,
        signature_path: str,
        placeholder_name: str,
        options: Optional[SignatureProcessingOptions] = None
    ) -> bool:
        """
        Integrate processed signature into document at placeholder location
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches
            
            # Load document
            doc = DocxDocument(document_path)
            signature_placeholder = f"${{{placeholder_name}}}"
            
            # Find and replace signature placeholder
            for paragraph in doc.paragraphs:
                if signature_placeholder in paragraph.text:
                    # Clear paragraph and add signature
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(signature_path, width=Inches(2.0))
                    break
            
            # Check tables for signature placeholders
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if signature_placeholder in paragraph.text:
                                paragraph.clear()
                                run = paragraph.add_run()
                                run.add_picture(signature_path, width=Inches(1.5))
            
            # Save document
            doc.save(document_path)
            
            return True
            
        except Exception as e:
            signature_logger.error(f"Failed to integrate signature: {e}")
            return False
    
    def calculate_optimal_signature_size(
        self,
        available_space: Tuple[int, int],
        signature_dimensions: Tuple[int, int]
    ) -> Tuple[int, int]:
        """
        Calculate optimal signature size for available document space
        """
        available_width, available_height = available_space
        sig_width, sig_height = signature_dimensions
        
        # Calculate aspect ratio
        aspect_ratio = sig_width / sig_height
        
        # Determine optimal size
        if available_width / available_height > aspect_ratio:
            # Height is limiting
            optimal_height = min(available_height, 100)  # Max 100px height
            optimal_width = int(optimal_height * aspect_ratio)
        else:
            # Width is limiting
            optimal_width = min(available_width, 300)  # Max 300px width
            optimal_height = int(optimal_width / aspect_ratio)
        
        return optimal_width, optimal_height

class SignatureCanvasService:
    """
    Service for managing signature canvas operations
    """
    
    def __init__(self):
        self.processor = AdvancedSignatureProcessor()
    
    async def process_canvas_signature(
        self,
        canvas_data: str,
        user_id: int,
        db: Session,
        document_id: Optional[int] = None,
        processing_options: Optional[SignatureProcessingOptions] = None
    ) -> Dict[str, Any]:
        """
        Complete signature processing workflow from canvas to database
        """
        try:
            # Process signature
            processed_bytes, processing_stats = await self.processor.process_signature_canvas(
                canvas_data, processing_options
            )
            
            # Save to database
            signature = await self.processor.save_signature_to_database(
                db, processed_bytes, user_id, document_id, 
                SignatureType.CANVAS, processing_stats
            )
            
            return {
                'signature_id': signature.id,
                'file_path': signature.file_path,
                'processing_stats': processing_stats,
                'status': 'success'
            }
            
        except Exception as e:
            signature_logger.error(f"Canvas signature processing failed: {e}")
            return {
                'status': 'error',
                'error': str(e)
            }
    
    async def get_signature_preview_url(
        self,
        signature_id: int,
        db: Session
    ) -> Optional[str]:
        """
        Get preview URL for signature
        """
        signature = db.query(Signature).filter(Signature.id == signature_id).first()
        
        if signature and signature.file_path:
            return f"/api/signatures/{signature_id}/preview"
        
        return None

# Global service instance
signature_canvas_service = SignatureCanvasService()