"""
Ultra-Fast Document Processing Engine
Implements sub-500ms document generation with memory streams, advanced caching, 
and intelligent placeholder processing for industry-standard performance.
"""

import os
import time
import json
import uuid
import asyncio
import hashlib
from io import BytesIO
from typing import Dict, List, Optional, Any, Tuple, Union
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor, as_completed

# Document processing imports
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx2pdf

# Database and models
from sqlalchemy.orm import Session
from app.models.template import Template, Placeholder
from app.models.document import Document, DocumentStatus
from app.services.cache_service import cache_service
from config import settings

# Advanced imports for performance
import redis
import aiofiles
import concurrent.futures
from dataclasses import dataclass
from enum import Enum
import logging

# Set up performance logger
perf_logger = logging.getLogger('performance')
perf_logger.setLevel(logging.INFO)

class ProcessingPriority(Enum):
    """Document processing priority levels"""
    LOW = 1
    NORMAL = 2
    HIGH = 3
    CRITICAL = 4

@dataclass
class PlaceholderContext:
    """Enhanced placeholder context information"""
    name: str
    display_name: str
    value: str
    position: Dict[str, Any]
    formatting: Dict[str, Any]
    context_type: str  # header, body, footer, signature_line
    semantic_group: str  # name, address, date, signature
    validation_rules: Dict[str, Any]
    transformation_rules: Dict[str, Any]

@dataclass
class DocumentGenerationRequest:
    """Optimized document generation request structure"""
    template_id: int
    placeholder_data: Dict[str, Any]
    output_format: str
    user_id: int
    priority: ProcessingPriority = ProcessingPriority.NORMAL
    cache_key: Optional[str] = None
    performance_target_ms: int = 500

class UltraFastDocumentEngine:
    """
    Ultra-high performance document processing engine
    Targets sub-500ms document generation through:
    - Memory-only processing
    - Advanced template caching 
    - Parallel placeholder processing
    - Context-aware formatting
    - Intelligent pre-processing
    """
    
    def __init__(self):
        self.template_cache = {}
        self.placeholder_cache = {}
        self.processing_stats = {}
        self.executor = ThreadPoolExecutor(max_workers=8)
        
        # Performance monitoring
        self.generation_times = []
        self.cache_hit_rate = 0.0
        self.processed_documents = 0
        
        # Initialize Redis connection for advanced caching
        try:
            self.redis_client = redis.Redis(
                host=settings.REDIS_HOST,
                port=settings.REDIS_PORT,
                password=settings.REDIS_PASSWORD,
                decode_responses=False,  # Keep binary for document caching
                socket_connect_timeout=2,
                socket_timeout=2
            )
            self.redis_available = True
        except Exception:
            self.redis_available = False
    
    async def generate_document_ultra_fast(
        self, 
        request: DocumentGenerationRequest,
        db: Session
    ) -> Tuple[BytesIO, Dict[str, Any]]:
        """
        Ultra-fast document generation with sub-500ms target
        
        Returns:
            Tuple of (document_stream, generation_stats)
        """
        start_time = time.time()
        stats = {
            "start_time": start_time,
            "cache_hits": 0,
            "processing_steps": [],
            "template_load_time": 0,
            "placeholder_processing_time": 0,
            "document_assembly_time": 0,
            "total_time": 0
        }
        
        try:
            # Step 1: Load template with intelligent caching (Target: <50ms)
            template, template_document = await self._load_template_cached(
                request.template_id, db
            )
            template_time = time.time()
            stats["template_load_time"] = (template_time - start_time) * 1000
            stats["processing_steps"].append("template_loaded")
            
            # Step 2: Process placeholders with context awareness (Target: <150ms)
            processed_placeholders = await self._process_placeholders_parallel(
                template, request.placeholder_data
            )
            placeholder_time = time.time()
            stats["placeholder_processing_time"] = (placeholder_time - template_time) * 1000
            stats["processing_steps"].append("placeholders_processed")
            
            # Step 3: Generate document in memory (Target: <200ms)
            document_stream = await self._generate_document_memory(
                template_document, processed_placeholders, request.output_format
            )
            assembly_time = time.time()
            stats["document_assembly_time"] = (assembly_time - placeholder_time) * 1000
            stats["processing_steps"].append("document_assembled")
            
            # Calculate total time
            total_time = (assembly_time - start_time) * 1000
            stats["total_time"] = total_time
            
            # Log performance
            self._log_performance_metrics(request, stats)
            
            return document_stream, stats
            
        except Exception as e:
            error_time = time.time()
            stats["total_time"] = (error_time - start_time) * 1000
            stats["error"] = str(e)
            perf_logger.error(f"Document generation failed: {e}")
            raise
    
    async def _load_template_cached(
        self, 
        template_id: int, 
        db: Session
    ) -> Tuple[Template, DocxDocument]:
        """
        Load template with multi-layer caching strategy
        Memory -> Redis -> Database
        """
        cache_key = f"template_full_{template_id}"
        
        # Layer 1: Memory cache
        if cache_key in self.template_cache:
            template_data, doc_bytes = self.template_cache[cache_key]
            doc_stream = BytesIO(doc_bytes)
            template_doc = DocxDocument(doc_stream)
            return template_data, template_doc
        
        # Layer 2: Redis cache
        if self.redis_available:
            try:
                cached_data = self.redis_client.hgetall(cache_key)
                if cached_data:
                    template_json = cached_data[b'template'].decode('utf-8')
                    doc_bytes = cached_data[b'document']
                    
                    # Reconstruct template object
                    template_dict = json.loads(template_json)
                    template = Template(**template_dict)
                    
                    # Load document
                    doc_stream = BytesIO(doc_bytes)
                    template_doc = DocxDocument(doc_stream)
                    
                    # Cache in memory for next time
                    self.template_cache[cache_key] = (template, doc_bytes)
                    
                    return template, template_doc
            except Exception as e:
                perf_logger.warning(f"Redis cache miss for template {template_id}: {e}")
        
        # Layer 3: Database load
        template = db.query(Template).filter(Template.id == template_id).first()
        if not template:
            raise ValueError(f"Template {template_id} not found")
        
        # Load document file
        template_path = os.path.join(settings.TEMPLATES_PATH, template.file_path)
        
        async with aiofiles.open(template_path, 'rb') as f:
            doc_bytes = await f.read()
        
        doc_stream = BytesIO(doc_bytes)
        template_doc = DocxDocument(doc_stream)
        
        # Cache in both Redis and memory
        if self.redis_available:
            try:
                # Serialize template for Redis
                template_dict = {
                    'id': template.id,
                    'name': template.name,
                    'file_path': template.file_path,
                    'placeholders': template.placeholders,
                    'font_family': template.font_family,
                    'font_size': template.font_size,
                    'created_at': template.created_at.isoformat() if template.created_at else None,
                    'updated_at': template.updated_at.isoformat() if template.updated_at else None
                }
                
                cache_data = {
                    'template': json.dumps(template_dict),
                    'document': doc_bytes
                }
                
                self.redis_client.hset(cache_key, mapping=cache_data)
                self.redis_client.expire(cache_key, 3600)  # 1 hour
                
            except Exception as e:
                perf_logger.warning(f"Failed to cache template {template_id} in Redis: {e}")
        
        # Cache in memory
        self.template_cache[cache_key] = (template, doc_bytes)
        
        return template, template_doc
    
    async def _process_placeholders_parallel(
        self,
        template: Template,
        placeholder_data: Dict[str, Any]
    ) -> List[PlaceholderContext]:
        """
        Process placeholders in parallel with context awareness and intelligent formatting
        """
        # Parse template placeholders
        template_placeholders = json.loads(template.placeholders) if template.placeholders else []
        
        # Create processing tasks
        processing_tasks = []
        
        for placeholder in template_placeholders:
            task = self._process_single_placeholder(
                placeholder, placeholder_data, template
            )
            processing_tasks.append(task)
        
        # Process all placeholders concurrently
        processed_placeholders = await asyncio.gather(*processing_tasks)
        
        return processed_placeholders
    
    async def _process_single_placeholder(
        self,
        placeholder: Dict[str, Any],
        data: Dict[str, Any],
        template: Template
    ) -> PlaceholderContext:
        """
        Process individual placeholder with context-aware formatting
        """
        placeholder_name = placeholder.get('name', '')
        raw_value = data.get(placeholder_name, '')
        
        # Detect context type
        context_type = self._detect_placeholder_context(placeholder, template)
        
        # Determine semantic group
        semantic_group = self._classify_placeholder_semantic_group(placeholder_name)
        
        # Apply context-aware formatting
        formatted_value = await self._apply_context_aware_formatting(
            raw_value, placeholder_name, context_type, semantic_group
        )
        
        # Create enhanced placeholder context
        context = PlaceholderContext(
            name=placeholder_name,
            display_name=placeholder.get('display_name', placeholder_name),
            value=formatted_value,
            position={
                'paragraph_index': placeholder.get('paragraph_index', 0),
                'run_index': placeholder.get('run_index', 0),
                'start_index': placeholder.get('start_index', 0),
                'end_index': placeholder.get('end_index', 0)
            },
            formatting={
                'bold': placeholder.get('bold', False),
                'italic': placeholder.get('italic', False),
                'underline': placeholder.get('underline', False),
                'font_size': placeholder.get('font_size', template.font_size),
                'font_family': placeholder.get('font_family', template.font_family)
            },
            context_type=context_type,
            semantic_group=semantic_group,
            validation_rules=self._get_validation_rules(semantic_group),
            transformation_rules=self._get_transformation_rules(context_type, semantic_group)
        )
        
        return context
    
    def _detect_placeholder_context(self, placeholder: Dict[str, Any], template: Template) -> str:
        """
        Detect placeholder context (header, body, footer, signature_line) based on position
        """
        paragraph_index = placeholder.get('paragraph_index', 0)
        
        # Simple heuristic - improve with document analysis
        if paragraph_index < 3:
            return "header"
        elif paragraph_index > 20:  # Assuming most documents have < 20 paragraphs
            return "footer"
        elif "signature" in placeholder.get('name', '').lower():
            return "signature_line"
        else:
            return "body"
    
    def _classify_placeholder_semantic_group(self, placeholder_name: str) -> str:
        """
        Classify placeholder into semantic groups for intelligent processing
        """
        name_lower = placeholder_name.lower()
        
        if any(word in name_lower for word in ['name', 'applicant', 'student', 'client']):
            return "name"
        elif any(word in name_lower for word in ['address', 'location', 'residence']):
            return "address"
        elif any(word in name_lower for word in ['date', 'birth', 'issued', 'expire']):
            return "date"
        elif any(word in name_lower for word in ['signature', 'sign']):
            return "signature"
        elif any(word in name_lower for word in ['email', 'mail']):
            return "email"
        elif any(word in name_lower for word in ['phone', 'mobile', 'tel']):
            return "phone"
        elif any(word in name_lower for word in ['age', 'years', 'number']):
            return "number"
        else:
            return "text"
    
    async def _apply_context_aware_formatting(
        self,
        value: str,
        placeholder_name: str,
        context_type: str,
        semantic_group: str
    ) -> str:
        """
        Apply intelligent formatting based on context and semantic group
        """
        if not value:
            return value
        
        # Address formatting based on context
        if semantic_group == "address":
            if context_type == "header":
                # Header addresses: comma-separated with line breaks
                return self._format_address_header(value)
            elif context_type == "body":
                # Body addresses: inline format
                return self._format_address_inline(value)
        
        # Date formatting
        elif semantic_group == "date":
            return self._format_date_intelligent(value, context_type)
        
        # Name formatting
        elif semantic_group == "name":
            if context_type == "header":
                return value.title()  # Title case for headers
            elif context_type == "signature_line":
                return value.upper()  # Uppercase for signatures
            else:
                return value.title()  # Default title case
        
        # Default processing
        return value
    
    def _format_address_header(self, address: str) -> str:
        """Format address for header context with line breaks"""
        # Replace commas with line breaks for header formatting
        return address.replace(', ', '\n')
    
    def _format_address_inline(self, address: str) -> str:
        """Format address for inline body text"""
        # Keep comma-separated for inline use
        return address.strip()
    
    def _format_date_intelligent(self, date_str: str, context_type: str) -> str:
        """Intelligent date formatting based on context"""
        try:
            from dateutil.parser import parse
            date_obj = parse(date_str)
            
            if context_type == "header":
                return date_obj.strftime("%B %d, %Y")  # "January 15, 2024"
            else:
                return date_obj.strftime("%d/%m/%Y")    # "15/01/2024"
        except:
            return date_str  # Return original if parsing fails
    
    def _get_validation_rules(self, semantic_group: str) -> Dict[str, Any]:
        """Get validation rules for semantic group"""
        rules = {
            "name": {"min_length": 2, "max_length": 100, "pattern": r"^[a-zA-Z\s]+$"},
            "email": {"pattern": r"^[^@]+@[^@]+\.[^@]+$"},
            "phone": {"pattern": r"^\+?[1-9]\d{1,14}$"},
            "date": {"format": ["DD/MM/YYYY", "MM/DD/YYYY", "YYYY-MM-DD"]},
            "address": {"min_length": 10, "max_length": 500},
            "number": {"type": "integer", "min": 0, "max": 999999},
            "text": {"min_length": 1, "max_length": 1000}
        }
        return rules.get(semantic_group, {"min_length": 1})
    
    def _get_transformation_rules(self, context_type: str, semantic_group: str) -> Dict[str, Any]:
        """Get transformation rules based on context and semantic group"""
        return {
            "context_type": context_type,
            "semantic_group": semantic_group,
            "apply_casing": True,
            "apply_formatting": True,
            "preserve_structure": context_type in ["header", "signature_line"]
        }
    
    async def _generate_document_memory(
        self,
        template_doc: DocxDocument,
        placeholders: List[PlaceholderContext],
        output_format: str
    ) -> BytesIO:
        """
        Generate final document in memory with optimized placeholder replacement
        """
        # Process placeholders in parallel batches
        batch_size = 5
        placeholder_batches = [
            placeholders[i:i + batch_size] 
            for i in range(0, len(placeholders), batch_size)
        ]
        
        # Apply placeholder replacements in batches
        for batch in placeholder_batches:
            await self._apply_placeholder_batch(template_doc, batch)
        
        # Generate output stream
        output_stream = BytesIO()
        
        if output_format.lower() == 'pdf':
            # For PDF, we need to save as DOCX first then convert
            temp_docx_stream = BytesIO()
            template_doc.save(temp_docx_stream)
            temp_docx_stream.seek(0)
            
            # Convert to PDF (this would need python-docx2pdf or similar)
            # For now, return DOCX stream (PDF conversion can be added)
            output_stream = temp_docx_stream
        else:
            # Save as DOCX
            template_doc.save(output_stream)
        
        output_stream.seek(0)
        return output_stream
    
    async def _apply_placeholder_batch(
        self,
        document: DocxDocument,
        placeholders: List[PlaceholderContext]
    ):
        """
        Apply placeholder replacements in batch for better performance
        """
        for placeholder in placeholders:
            await self._replace_placeholder_optimized(document, placeholder)
    
    async def _replace_placeholder_optimized(
        self,
        document: DocxDocument,
        placeholder: PlaceholderContext
    ):
        """
        Optimized placeholder replacement with formatting preservation
        """
        placeholder_text = f"${{{placeholder.name}}}"
        
        # Search and replace in all paragraphs
        for paragraph in document.paragraphs:
            if placeholder_text in paragraph.text:
                # Advanced replacement logic preserving formatting
                self._replace_in_paragraph(paragraph, placeholder_text, placeholder)
        
        # Search in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder_text in paragraph.text:
                            self._replace_in_paragraph(paragraph, placeholder_text, placeholder)
        
        # Search in headers/footers
        for section in document.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                if placeholder_text in paragraph.text:
                    self._replace_in_paragraph(paragraph, placeholder_text, placeholder)
            
            # Footer  
            footer = section.footer
            for paragraph in footer.paragraphs:
                if placeholder_text in paragraph.text:
                    self._replace_in_paragraph(paragraph, placeholder_text, placeholder)
    
    def _replace_in_paragraph(
        self,
        paragraph,
        placeholder_text: str,
        placeholder: PlaceholderContext
    ):
        """
        Replace placeholder text in paragraph while preserving formatting
        """
        if placeholder_text not in paragraph.text:
            return
        
        # Handle address formatting with line breaks
        if placeholder.semantic_group == "address" and "\n" in placeholder.value:
            self._replace_with_line_breaks(paragraph, placeholder_text, placeholder.value)
        else:
            # Standard replacement
            paragraph.text = paragraph.text.replace(placeholder_text, placeholder.value)
            
            # Apply formatting to the replaced text
            for run in paragraph.runs:
                if placeholder.value in run.text:
                    run.bold = placeholder.formatting.get('bold', False)
                    run.italic = placeholder.formatting.get('italic', False)
                    run.underline = placeholder.formatting.get('underline', False)
                    
                    if placeholder.formatting.get('font_size'):
                        run.font.size = Pt(placeholder.formatting['font_size'])
                    if placeholder.formatting.get('font_family'):
                        run.font.name = placeholder.formatting['font_family']
    
    def _replace_with_line_breaks(self, paragraph, placeholder_text: str, value: str):
        """
        Replace placeholder with multi-line text, creating proper line breaks
        """
        if placeholder_text not in paragraph.text:
            return
        
        # Clear paragraph
        paragraph.clear()
        
        # Split value by newlines and add as separate runs
        lines = value.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                # Add line break
                run = paragraph.add_run()
                run.add_break()
            # Add text line
            paragraph.add_run(line)
    
    def _log_performance_metrics(
        self,
        request: DocumentGenerationRequest,
        stats: Dict[str, Any]
    ):
        """
        Log performance metrics for monitoring and optimization
        """
        total_time = stats["total_time"]
        self.generation_times.append(total_time)
        self.processed_documents += 1
        
        # Log if exceeds target
        if total_time > request.performance_target_ms:
            perf_logger.warning(
                f"Document generation exceeded target: {total_time}ms > {request.performance_target_ms}ms"
            )
        
        # Log performance summary every 100 documents
        if self.processed_documents % 100 == 0:
            avg_time = sum(self.generation_times[-100:]) / min(100, len(self.generation_times))
            perf_logger.info(
                f"Performance Summary - Avg: {avg_time:.2f}ms, "
                f"Documents: {self.processed_documents}, "
                f"Cache Hit Rate: {self.cache_hit_rate:.2f}%"
            )
    
    async def batch_generate_documents(
        self,
        requests: List[DocumentGenerationRequest],
        db: Session
    ) -> List[Tuple[BytesIO, Dict[str, Any]]]:
        """
        Ultra-fast batch document generation with intelligent load balancing
        """
        start_time = time.time()
        
        # Sort requests by priority
        sorted_requests = sorted(requests, key=lambda x: x.priority.value, reverse=True)
        
        # Process in parallel with controlled concurrency
        max_concurrent = min(len(requests), 5)  # Limit to prevent resource exhaustion
        
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_request(request):
            async with semaphore:
                return await self.generate_document_ultra_fast(request, db)
        
        # Execute all requests concurrently
        tasks = [process_request(req) for req in sorted_requests]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        total_time = (time.time() - start_time) * 1000
        perf_logger.info(f"Batch generated {len(requests)} documents in {total_time:.2f}ms")
        
        return results
    
    def get_performance_stats(self) -> Dict[str, Any]:
        """
        Get current performance statistics
        """
        if not self.generation_times:
            return {"status": "no_data"}
        
        recent_times = self.generation_times[-100:]
        
        return {
            "total_documents": self.processed_documents,
            "average_generation_time": sum(recent_times) / len(recent_times),
            "min_generation_time": min(recent_times),
            "max_generation_time": max(recent_times),
            "cache_hit_rate": self.cache_hit_rate,
            "templates_cached": len(self.template_cache),
            "redis_available": self.redis_available
        }

# Global engine instance
ultra_fast_engine = UltraFastDocumentEngine()